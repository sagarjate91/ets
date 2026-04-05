"""
Convert Capstone_Project_Report.html → Capstone_Project_Report.docx
Preserves: headings, paragraphs, tables, code blocks, images (snapshots)
"""

import re, os, base64
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(r"C:\Users\sagar\Documents\23 March 2026\ets")
HTML_FILE = BASE_DIR / "Capstone_Project_Report.html"
OUT_FILE  = BASE_DIR / "Capstone_Project_Report.docx"
SNAP_DIR  = BASE_DIR / "snapshots"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "4472C4")
    pb.append(bot)
    pPr.append(pb)
    return p

def apply_run_style(run, bold=False, italic=False, color=None, size=None, font_name=None):
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if size:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name

def para_align(p, align):
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# Normal style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# Heading styles
_heading_defs = {
    "Heading 1": ("Calibri", 18, "0B3D91", True),
    "Heading 2": ("Calibri", 15, "1155BB", True),
    "Heading 3": ("Calibri", 13, "1155BB", False),
    "Heading 4": ("Calibri", 11, "333333", False),
}
for name, (font, size, color, bold) in _heading_defs.items():
    s = doc.styles[name]
    s.font.name  = font
    s.font.size  = Pt(size)
    s.font.color.rgb = RGBColor(*bytes.fromhex(color))
    s.font.bold  = bold

# Code style
try:
    code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
except:
    code_style = doc.styles["CodeBlock"]
code_style.font.name = "Courier New"
code_style.font.size = Pt(8)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after  = Pt(2)

# ── parse HTML ────────────────────────────────────────────────────────────────

html_text = HTML_FILE.read_text(encoding="utf-8")
soup = BeautifulSoup(html_text, "lxml")

# Remove <head> / <script> / <style> / nav
for tag in soup.find_all(["head", "script", "style", "nav"]):
    tag.decompose()

body = soup.find("body") or soup

# ── cover page ────────────────────────────────────────────────────────────────

def make_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("CAPSTONE PROJECT REPORT")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)
    run.font.name = "Calibri"

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Employee Tracking System (ETS)")
    r.bold = True; r.font.size = Pt(16); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0x11, 0x55, 0xBB)

    doc.add_paragraph()
    info = [
        ("Technology:", "Java 17 · Spring Boot · Spring Security · Thymeleaf · MySQL / H2"),
        ("Prepared by:", "Sagar"),
        ("Date:",        "April 2026"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "  ")
        r1.bold = True; r1.font.size = Pt(12); r1.font.name = "Calibri"
        r2 = p.add_run(value)
        r2.font.size = Pt(12); r2.font.name = "Calibri"

    doc.add_page_break()

make_cover(doc)

# ── recursive content walker ──────────────────────────────────────────────────

def inner_text(node):
    return node.get_text(separator=" ", strip=True)

def process_node(node, doc, list_level=0):
    if isinstance(node, str):
        return

    tag = node.name if node.name else ""

    # ── headings ────────────────────────────────────────────────────────────
    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        level = min(level, 4)  # cap at Heading 4
        txt = inner_text(node).strip()
        if not txt:
            return
        # Page break before h2
        if level == 2:
            doc.add_page_break()
        p = doc.add_heading(txt, level=level)
        return

    # ── paragraphs ──────────────────────────────────────────────────────────
    if tag == "p":
        txt = inner_text(node).strip()
        if not txt:
            return
        p = doc.add_paragraph()
        run = p.add_run(txt)
        run.font.size = Pt(11)
        return

    # ── lists ───────────────────────────────────────────────────────────────
    if tag in ("ul", "ol"):
        for li in node.find_all("li", recursive=False):
            txt = inner_text(li).strip()
            if txt:
                p = doc.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
                run = p.add_run(txt)
                run.font.size = Pt(11)
        return

    # ── code blocks ─────────────────────────────────────────────────────────
    if tag == "pre":
        code_node = node.find("code")
        raw = (code_node or node).get_text()
        lines = raw.split("\n")
        # Trim very long blocks (>300 lines → keep first 300 + note)
        MAX = 300
        truncated = len(lines) > MAX
        if truncated:
            lines = lines[:MAX]
        for line in lines:
            p = doc.add_paragraph(style="CodeBlock")
            p.add_run(line if line else " ")
        if truncated:
            p = doc.add_paragraph(style="CodeBlock")
            r = p.add_run(f"  ... [ {len(raw.split(chr(10))) - MAX} more lines truncated for readability ]")
            r.italic = True
        return

    # ── inline code ─────────────────────────────────────────────────────────
    if tag == "code":
        txt = node.get_text(strip=True)
        if txt:
            p = doc.add_paragraph()
            r = p.add_run(txt)
            r.font.name = "Courier New"
            r.font.size = Pt(9)
        return

    # ── images ──────────────────────────────────────────────────────────────
    if tag == "img":
        src = node.get("src", "")
        alt = node.get("alt", "Image")
        img_path = None
        if src.startswith("snapshots/"):
            img_path = BASE_DIR / src
        elif src.startswith("data:image"):
            # base64 embedded
            try:
                _, data = src.split(",", 1)
                import io
                img_bytes = base64.b64decode(data)
                img_path = io.BytesIO(img_bytes)
            except Exception:
                img_path = None

        if img_path:
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path) if isinstance(img_path, Path) else img_path,
                                width=Inches(5.5))
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
            except Exception as e:
                p = doc.add_paragraph(f"[Image: {alt}]")
        else:
            doc.add_paragraph(f"[Image: {alt}]")
        return

    # ── tables ──────────────────────────────────────────────────────────────
    if tag == "table":
        rows_el = node.find_all("tr")
        if not rows_el:
            return
        col_count = max(
            len(r.find_all(["td", "th"])) for r in rows_el
        )
        if col_count == 0:
            return
        tbl = doc.add_table(rows=0, cols=col_count)
        tbl.style = "Table Grid"
        for i, tr in enumerate(rows_el):
            cells_el = tr.find_all(["td", "th"])
            row = tbl.add_row()
            for j, cell_el in enumerate(cells_el[:col_count]):
                cell = row.cells[j]
                txt  = inner_text(cell_el).strip()
                cell.text = txt
                if cell_el.name == "th" or i == 0:
                    set_cell_bg(cell, "D6E4F7")
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                else:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)
        doc.add_paragraph()
        return

    # ── horizontal rule ─────────────────────────────────────────────────────
    if tag == "hr":
        add_horizontal_rule(doc)
        return

    # ── page break ──────────────────────────────────────────────────────────
    if tag in ("section",) and "page-break" in node.get("class", []):
        # handled by h2 above; just process children
        pass

    # ── skip nav / toc ──────────────────────────────────────────────────────
    if node.get("id") in ("toc",):
        return
    if "toc" in node.get("class", []):
        return

    # ── div / section / article / span / generic ────────────────────────────
    for child in node.children:
        process_node(child, doc, list_level)

# ── walk body ─────────────────────────────────────────────────────────────────

print("Processing HTML…")
for child in body.children:
    process_node(child, doc)

# ── save ─────────────────────────────────────────────────────────────────────

doc.save(str(OUT_FILE))
size = OUT_FILE.stat().st_size
print(f"\nSaved: {OUT_FILE}")
print(f"Size : {size:,} bytes")
